"""
generate_doc.py — Generate a deliverable by reading a template's sections and
their instruction text, then filling each section with project information.

HOW IT WORKS
------------
A template is the user's own document. Each section usually has a heading plus
some instruction or guidance text describing what belongs there (for example,
under "System Overview": "Describe the system, its purpose, and where it runs.").
This script and the generate-doc skill work together in three steps:

  1. Outline   — extract every section: its heading and its instruction text.
                 python scripts/generate_doc.py <template> --outline
  2. Compose   — the generate-doc skill reads the outline and the project
                 context (MASTER_CONTEXT.md), then writes, for each section, the
                 real content that satisfies the instruction. Where the project
                 does not have the information, it writes a [CONFIRM: ...] marker
                 instead. The skill saves this as a JSON map: {heading: content}.
  3. Fill      — replace each section's instruction text with the composed
                 content, keeping the template's headings, order, and formatting.
                 python scripts/generate_doc.py <template> --fill content.json

Running with no flag copies the template to deliverables/in-progress/ as a
starting draft and prints the outline so a draft always exists.

INPUTS
------
- template: a filename in templates/, or an alias under deliverables in
  workbench.config.yaml.
- MASTER_CONTEXT.md: project data (read by the skill, checked here).
- workbench.config.yaml: aliases, naming pattern.

DEPENDENCIES
------------
python-docx, openpyxl, PyYAML  (see requirements.txt). Markdown templates need
no extra packages; Word templates need python-docx.
"""

import argparse
import json
import re
import sys
from datetime import date
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

REPO_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = REPO_ROOT / "templates"
SOPS_DIR = REPO_ROOT / "sops"
OUTPUT_DIR = REPO_ROOT / "deliverables" / "in-progress"
CONTEXT_FILE = REPO_ROOT / "MASTER_CONTEXT.md"
STATUS_FILE = REPO_ROOT / "DELIVERABLE_STATUS.md"
CONFIG_FILE = REPO_ROOT / "workbench.config.yaml"

DEFAULT_NAMING = "{system}_{doc}_v{version}_draft"
DEFAULT_VERSION = "1.0"


def load_config():
    """Load workbench.config.yaml. Return {} if missing or PyYAML unavailable."""
    if not CONFIG_FILE.is_file():
        return {}
    try:
        import yaml
    except ImportError:
        return {}
    try:
        return yaml.safe_load(CONFIG_FILE.read_text(encoding="utf-8")) or {}
    except Exception as exc:
        print("Warning: could not parse workbench.config.yaml: {}".format(exc))
        return {}


def resolve_template(arg, config):
    """Resolve a template argument to a path in templates/.

    Order: config alias -> exact filename -> case-insensitive fuzzy match.
    """
    if not TEMPLATES_DIR.is_dir():
        return None
    candidates = [p for p in TEMPLATES_DIR.iterdir()
                  if p.is_file() and p.name not in (".gitkeep", "README.md")]
    for entry in config.get("deliverables") or []:
        if isinstance(entry, dict) and entry.get("alias", "").lower() == arg.lower():
            mapped = TEMPLATES_DIR / entry.get("template", "")
            if mapped.is_file():
                return mapped
    for path in candidates:
        if path.name.lower() == arg.lower():
            return path
    for path in candidates:
        if arg.lower() in path.stem.lower():
            return path
    return None


def system_name_from_context():
    if not CONTEXT_FILE.is_file():
        return None
    text = CONTEXT_FILE.read_text(encoding="utf-8", errors="replace")
    match = re.search(r"\*\*System Name:\*\*\s*(.+)", text)
    if match:
        name = match.group(1).strip()
        if name and "populated by" not in name and "CONFIRM" not in name:
            return name
    return None


def output_name(config, system, doc, suffix):
    naming = (config.get("naming") or {})
    pattern = naming.get("draft_pattern", DEFAULT_NAMING)
    version = naming.get("default_version", DEFAULT_VERSION)
    safe = lambda s: re.sub(r"[^A-Za-z0-9._-]+", "_", s).strip("_")
    stem = pattern.format(system=safe(system), doc=safe(doc),
                          version=version, date=date.today().isoformat())
    return stem + suffix


# --- Outline extraction -----------------------------------------------------

def _heading_level_from_style(style_name):
    match = re.search(r"(\d+)", style_name or "")
    return int(match.group(1)) if match else 1


def extract_outline_md(text):
    sections = []
    current = None
    for line in text.splitlines():
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            if current:
                current["instruction"] = current["instruction"].strip("\n")
                sections.append(current)
            current = {"level": len(m.group(1)), "heading": m.group(2).strip(), "instruction": ""}
        elif current is not None:
            current["instruction"] += line + "\n"
    if current:
        current["instruction"] = current["instruction"].strip("\n")
        sections.append(current)
    return sections


def extract_outline_docx(path):
    from docx import Document
    doc = Document(str(path))
    sections = []
    current = None
    for para in doc.paragraphs:
        style = (para.style.name if para.style else "") or ""
        is_heading = style.startswith("Heading") or style == "Title"
        if is_heading and para.text.strip():
            if current:
                current["instruction"] = current["instruction"].strip("\n")
                sections.append(current)
            current = {"level": _heading_level_from_style(style),
                       "heading": para.text.strip(), "instruction": ""}
        elif current is not None and para.text.strip():
            current["instruction"] += para.text + "\n"
    if current:
        current["instruction"] = current["instruction"].strip("\n")
        sections.append(current)
    return sections


def extract_outline(path):
    if path.suffix.lower() == ".docx":
        return extract_outline_docx(path)
    return extract_outline_md(path.read_text(encoding="utf-8", errors="replace"))


# --- Filling ----------------------------------------------------------------

def fill_md(text, content_map):
    """Rebuild a markdown template, replacing each section's body with mapped content."""
    out = []
    current_heading = None
    buffer_emitted = False
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            out.append(line)
            current_heading = m.group(2).strip()
            if current_heading in content_map:
                out.append("")
                out.append(content_map[current_heading].rstrip())
                # skip original body lines until the next heading
                i += 1
                while i < len(lines) and not re.match(r"^#{1,6}\s+", lines[i]):
                    i += 1
                continue
        else:
            out.append(line)
        i += 1
    return "\n".join(out) + "\n"


def fill_docx(path, content_map, dest):
    from docx import Document
    doc = Document(str(path))
    current_heading = None
    wrote_for_heading = set()
    for para in doc.paragraphs:
        style = (para.style.name if para.style else "") or ""
        is_heading = style.startswith("Heading") or style == "Title"
        if is_heading and para.text.strip():
            current_heading = para.text.strip()
            continue
        if current_heading in content_map and para.text.strip():
            if current_heading not in wrote_for_heading:
                _set_paragraph_text(para, content_map[current_heading])
                wrote_for_heading.add(current_heading)
            else:
                _set_paragraph_text(para, "")
    doc.save(str(dest))


def _set_paragraph_text(paragraph, new_text):
    """Replace a paragraph's text while keeping its style. Clears extra runs."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


# --- Modes ------------------------------------------------------------------

def do_outline(template):
    outline = extract_outline(template)
    print(json.dumps(outline, indent=2, ensure_ascii=False))


def do_fill(template, config, system, doc_label, fill_path):
    try:
        content_map = json.loads(Path(fill_path).read_text(encoding="utf-8"))
    except Exception as exc:
        print("Could not read the content map '{}': {}".format(fill_path, exc))
        return 1
    if not isinstance(content_map, dict):
        print("The content map must be a JSON object of {\"heading\": \"content\"}.")
        return 1

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    suffix = template.suffix.lower()
    dest = OUTPUT_DIR / output_name(config, system, doc_label, suffix if suffix in (".docx", ".md") else ".md")
    try:
        if suffix == ".docx":
            fill_docx(template, content_map, dest)
        else:
            filled = fill_md(template.read_text(encoding="utf-8", errors="replace"), content_map)
            dest.write_text(filled, encoding="utf-8")
    except ImportError:
        print("python-docx is required for .docx templates. Run setup first.")
        return 1
    except Exception as exc:
        print("Could not fill the template: {}".format(exc))
        return 1

    filled_text = dest.read_text(encoding="utf-8", errors="replace") if suffix != ".docx" else ""
    confirms = filled_text.count("[CONFIRM:")
    print("Filled draft created: {}".format(dest.relative_to(REPO_ROOT).as_posix()))
    if suffix != ".docx":
        print("{} section(s) still need a confirmed value.".format(confirms))
    return 0


def do_copy(template, config, system, doc_label):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    suffix = template.suffix.lower()
    dest = OUTPUT_DIR / output_name(config, system, doc_label, suffix if suffix in (".docx", ".md") else ".md")
    try:
        if suffix == ".docx":
            from docx import Document
            Document(str(template)).save(str(dest))
        else:
            dest.write_text(template.read_text(encoding="utf-8", errors="replace"), encoding="utf-8")
    except ImportError:
        print("python-docx is required for .docx templates. Run setup first.")
        return 1
    except Exception as exc:
        print("Could not create the draft: {}".format(exc))
        return 1
    print("Draft created: {}".format(dest.relative_to(REPO_ROOT).as_posix()))
    print("Sections found in this template:")
    for sec in extract_outline(template):
        print("  - {}".format(sec["heading"]))
    print("Next: the generate-doc skill fills each section from your project context.")
    return 0


def main(argv=None):
    parser = argparse.ArgumentParser(description="Generate a deliverable from a template.")
    parser.add_argument("template", help="Template filename or config alias.")
    parser.add_argument("--outline", action="store_true",
                        help="Print the template's sections and instruction text as JSON.")
    parser.add_argument("--fill", metavar="MAP.json", default=None,
                        help="Fill sections from a JSON map {heading: content} and write the draft.")
    parser.add_argument("--system", default=None, help="System name (defaults to MASTER_CONTEXT).")
    parser.add_argument("--doc", default=None, help="Document label for the filename.")
    args = parser.parse_args(argv)

    config = load_config()
    template = resolve_template(args.template, config)
    if template is None:
        available = [p.name for p in TEMPLATES_DIR.iterdir()
                     if p.is_file() and p.name not in (".gitkeep", "README.md")] if TEMPLATES_DIR.is_dir() else []
        print("No template matched '{}'.".format(args.template))
        if available:
            print("Available templates: " + ", ".join(available))
        else:
            print("templates/ has no templates yet. Add one, then re-run.")
        return 1

    if args.outline:
        do_outline(template)
        return 0

    # Context must be built before composing or copying a real draft.
    if not CONTEXT_FILE.is_file() or "[populated by /build-context]" in CONTEXT_FILE.read_text(encoding="utf-8", errors="replace"):
        print("MASTER_CONTEXT.md is not populated. Run /build-context first.")
        return 1

    system = args.system or system_name_from_context() or "System"
    doc_label = args.doc or args.template

    if args.fill:
        return do_fill(template, config, system, doc_label, args.fill)
    return do_copy(template, config, system, doc_label)


if __name__ == "__main__":
    sys.exit(main())
